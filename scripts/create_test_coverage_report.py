from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\basha\SE\nubl")
OUT = ROOT / "docs" / "test-coverage-investigation-report.docx"


COLORS = {
    "navy": "1F4E79",
    "blue": "2F75B5",
    "light_blue": "D9EAF7",
    "pale_blue": "EEF6FC",
    "green": "548235",
    "light_green": "E2F0D9",
    "gray": "6B7280",
    "light_gray": "F3F6F8",
    "border": "B7C9D6",
    "white": "FFFFFF",
    "red": "C00000",
}


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = COLORS["border"], size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, color=COLORS["white"], size=8)
        shade_cell(hdr[idx], COLORS["navy"])
        set_cell_border(hdr[idx])
        set_cell_margins(hdr[idx])
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=8)
            if len(value) < 18 and idx != 0:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor.from_string(COLORS["navy"] if level == 1 else COLORS["blue"])
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Cascadia Mono"
    run.font.size = Pt(7.2)
    run.font.color.rgb = RGBColor.from_string("1F2937")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), COLORS["light_gray"])
    p_pr.append(shd)


def add_callout(doc: Document, title: str, body: str, fill: str = COLORS["pale_blue"]) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, color=COLORS["border"])
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.1
    r2 = p2.add_run(body)
    r2.font.name = "Aptos"
    r2.font.size = Pt(9)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    for style_name, size in [("Title", 24), ("Subtitle", 12), ("Heading 1", 15), ("Heading 2", 11)]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        if "Heading" in style_name:
            style.font.color.rgb = RGBColor.from_string(COLORS["navy"])


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    r = p.add_run("Test Coverage Investigation Report")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(COLORS["navy"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NUBL Laravel Application")
    r.font.name = "Aptos Display"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(COLORS["gray"])

    doc.add_paragraph()
    add_callout(
        doc,
        "Final Judgment",
        "The PHP/Laravel application test coverage is 93.3% for the application source under app/. "
        "This judgment is supported by a successful PHPUnit coverage run using Xdebug, and by separate "
        "suite validation showing all Unit and Feature tests passing.",
        fill=COLORS["light_green"],
    )

    metadata = [
        ["Project path", r"C:\Users\basha\SE\nubl"],
        ["Report timestamp", "2026-05-07 22:00:19 +03:00"],
        ["Git commit", "ed13a87"],
        ["Working tree", "Clean at the time of evidence capture"],
        ["PHP runtime", "PHP 8.3.16 CLI"],
        ["Coverage driver", "Xdebug 3.5.1"],
    ]
    add_table(doc, ["Field", "Evidence"], metadata, widths=[1.8, 5.6])
    doc.add_page_break()


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. Executive Summary")
    add_body(
        doc,
        "The project was investigated as a Laravel/PHP application with PHPUnit configured for two PHP test suites: "
        "Unit and Feature. Xdebug was available in the PHP runtime, so line coverage could be measured directly. "
        "The authoritative coverage command completed successfully and reported a total coverage of 93.3%."
    )
    add_table(
        doc,
        ["Metric", "Result", "Evidence"],
        [
            ["Overall PHP coverage", "93.3%", "php artisan test --coverage reported: Total: 93.3%"],
            ["PHP test result", "784 passed", "php artisan test --compact"],
            ["Assertions", "2,658", "php artisan test --compact"],
            ["Unit tests", "243 passed", "php artisan test --testsuite=Unit --compact"],
            ["Feature tests", "541 passed", "php artisan test --testsuite=Feature --compact"],
            ["E2E tests detected", "5 files", "tests/e2e plus Playwright scripts in package.json"],
        ],
        widths=[2.0, 1.45, 4.1],
    )
    add_callout(
        doc,
        "Scope clarification",
        "The 93.3% value is PHPUnit coverage for Laravel application code included by phpunit.xml under app/. "
        "It does not include Playwright E2E coverage or front-end JavaScript coverage, because those are executed "
        "through separate npm scripts and are outside the PHPUnit coverage report.",
    )

    add_heading(doc, "2. Evidence Chain")
    rows = [
        ["1", "Identify framework and test tooling", "composer.json, artisan, phpunit.xml", "Laravel 12 project using PHPUnit"],
        ["2", "Confirm test suite structure", "phpunit.xml and tests folder", "Two PHPUnit suites: Unit and Feature"],
        ["3", "Confirm coverage capability", "php -v and php -m", "Xdebug 3.5.1 installed and xdebug module loaded"],
        ["4", "Run authoritative coverage command", "$env:XDEBUG_MODE='coverage'; php artisan test --coverage", "All tests passed; Total: 93.3%"],
        ["5", "Validate test counts by suite", "php artisan test --testsuite=Unit/Feature --compact", "243 Unit tests and 541 Feature tests passed"],
        ["6", "Check for other test layers", "package.json and tests/e2e", "5 Playwright E2E files exist but are not part of PHPUnit coverage"],
    ]
    add_table(doc, ["Step", "Investigation action", "Evidence source", "Conclusion"], rows, widths=[0.45, 2.25, 2.35, 2.45])

    add_heading(doc, "3. Commands and Results", level=1)
    command_rows = [
        [
            "Get-ChildItem -Force",
            "Repository root contains app, config, database, routes, tests, vendor, composer.json, package.json, phpunit.xml, and artisan.",
            "Confirmed a Laravel-style project root and local dependencies.",
        ],
        [
            "rg --files -g ...",
            "PowerShell returned: Program 'rg.exe' failed to run: Access is denied.",
            "Fallback used PowerShell file listing and Select-String; this did not affect the final coverage measurement.",
        ],
        [
            "Get-Content composer.json",
            "require includes laravel/framework ^12.0; require-dev includes phpunit/phpunit ^11.5.3; composer test script runs php artisan test.",
            "Confirmed PHP/Laravel test stack.",
        ],
        [
            "Get-Content package.json",
            "Scripts include test:e2e, test:e2e:php-only, test:e2e:ui, and test:e2e:headed using Playwright.",
            "Confirmed a separate browser E2E layer exists.",
        ],
        [
            "Get-Content phpunit.xml",
            "testsuites define Unit => tests/Unit and Feature => tests/Feature. source include is app. APP_ENV=testing and sqlite :memory: are configured.",
            "Confirmed what the coverage percentage is measuring.",
        ],
        [
            "Get-ChildItem -Path tests -Recurse -File",
            "Detected Unit, Feature, and e2e test directories with many PHP tests and 5 Playwright specs.",
            "Confirmed test organization.",
        ],
        [
            "php -v",
            "PHP 8.3.16 CLI with Xdebug v3.5.1.",
            "Coverage can be generated locally.",
        ],
        [
            "php -m",
            "Module list includes xdebug.",
            "Confirmed coverage driver is loaded.",
        ],
        [
            "Get-ChildItem tests\\Unit / tests\\Feature / tests\\e2e counts",
            "46 Unit PHP files, 70 Feature PHP files, and 5 E2E JavaScript files.",
            "Confirmed file-level distribution.",
        ],
        [
            "$env:XDEBUG_MODE='coverage'; php artisan test --coverage",
            "Command exited successfully after 855.7 seconds. Coverage report ended with: Total: 93.3%.",
            "Authoritative basis for final coverage judgment.",
        ],
        [
            "php artisan test --compact",
            "Tests: 784 passed (2658 assertions). Duration: 349.24s.",
            "Confirmed all PHP tests pass without coverage overhead.",
        ],
        [
            "php artisan test --testsuite=Unit --compact",
            "Tests: 243 passed (719 assertions). Duration: 34.41s.",
            "Confirmed Unit suite result.",
        ],
        [
            "php artisan test --testsuite=Feature --compact",
            "Tests: 541 passed (1939 assertions). Duration: 216.73s.",
            "Confirmed Feature suite result.",
        ],
        [
            "Get-Date / git rev-parse --short HEAD / git status --short",
            "Timestamp: 2026-05-07 22:00:19 +03:00. Commit: ed13a87. git status --short returned no output.",
            "Captured reproducibility metadata.",
        ],
    ]
    add_table(doc, ["Command", "Observed result", "Interpretation"], command_rows, widths=[2.35, 3.2, 2.0])

    add_heading(doc, "4. Key Configuration Evidence")
    add_body(doc, "The following configuration snippets explain why the reported percentage should be interpreted as Laravel/PHP application coverage.")
    add_code_block(
        doc,
        'phpunit.xml\n'
        '<testsuite name="Unit">\n'
        '    <directory>tests/Unit</directory>\n'
        '</testsuite>\n'
        '<testsuite name="Feature">\n'
        '    <directory>tests/Feature</directory>\n'
        '</testsuite>\n'
        '<source>\n'
        '    <include>\n'
        '        <directory>app</directory>\n'
        '    </include>\n'
        '</source>'
    )
    add_code_block(
        doc,
        'package.json\n'
        '"test:e2e": "playwright test --config=config/playwright/playwright.config.js",\n'
        '"test:e2e:php-only": "playwright test --config=config/playwright/playwright.php-only.config.js",\n'
        '"test:e2e:ui": "playwright test --config=config/playwright/playwright.config.js --ui",\n'
        '"test:e2e:headed": "playwright test --config=config/playwright/playwright.config.js --headed"'
    )

    add_heading(doc, "5. Coverage Result Evidence")
    add_body(doc, "The coverage command generated a long per-file coverage report. The decision-critical output is summarized below.")
    add_code_block(
        doc,
        "$env:XDEBUG_MODE='coverage'; php artisan test --coverage\n\n"
        "Exit code: 0\n"
        "Representative result: all Unit and Feature suites passed.\n"
        "Coverage report tail:\n"
        "  View/Components\\RegisterLayout .............................................. 100.0%\n"
        "  View/Composers\\SidebarComposer ................................. 14..19, 45 / 82.4%\n"
        "  ------------------------------------------------------------------------------\n"
        "                                                                  Total: 93.3 %"
    )
    add_body(
        doc,
        "Several individual app components reached 100.0% coverage, while lower coverage was observed in selected testing-only controllers, middleware, and notification/controller paths. These per-file details do not change the aggregate conclusion; they explain how the total was formed."
    )

    add_heading(doc, "6. Unit vs Feature Classification")
    add_table(
        doc,
        ["Category", "Directory", "Files detected", "Tests executed", "Assertions", "Result"],
        [
            ["Unit", "tests/Unit", "46 PHP files", "243", "719", "Passed"],
            ["Feature", "tests/Feature", "70 PHP files", "541", "1,939", "Passed"],
            ["E2E", "tests/e2e", "5 JS files", "Not run in PHPUnit", "N/A", "Separate Playwright layer"],
        ],
        widths=[0.8, 1.35, 1.15, 1.2, 1.0, 2.0],
    )
    add_body(
        doc,
        "The PHP test suite is therefore not only Unit tests or only Feature tests. It contains both, with the majority of executed PHP tests classified as Feature tests."
    )

    add_heading(doc, "7. Professional Judgment")
    add_callout(
        doc,
        "Conclusion",
        "The project has strong PHP application test coverage at 93.3%, measured against the Laravel app/ source tree. "
        "Both Unit and Feature suites are present and passing. The result is credible because the coverage driver was confirmed, "
        "the coverage command succeeded, and independent compact suite runs reproduced the passing test counts.",
        fill=COLORS["light_green"],
    )
    add_body(
        doc,
        "For reporting precision, state the result as: 'PHPUnit coverage for the Laravel application source under app/ is 93.3%. The PHPUnit suite consists of 243 Unit tests and 541 Feature tests, all passing. A separate Playwright E2E test layer exists but is not included in this coverage percentage.'"
    )

    add_heading(doc, "8. Limitations and Recommendations")
    add_table(
        doc,
        ["Point", "Implication", "Recommended wording or action"],
        [
            [
                "Coverage scope is app/",
                "Routes, views, config, migrations, and front-end assets are not counted unless executed through app code.",
                "Report the metric as PHP/Laravel app coverage, not total repository coverage.",
            ],
            [
                "Playwright E2E tests are separate",
                "They validate browser workflows but did not contribute to the 93.3% PHPUnit coverage number.",
                "Mention E2E coverage separately if a browser coverage tool is later configured.",
            ],
            [
                "Long coverage run",
                "Coverage with Xdebug took about 14 minutes on this machine.",
                "Use compact test runs for quick validation and coverage runs for formal evidence.",
            ],
            [
                "Some low-coverage files remain",
                "Testing-only controllers/middleware and selected workflow branches had lower percentages.",
                "If the target is above 95%, prioritize files below 80% and high-risk branches first.",
            ],
        ],
        widths=[1.5, 2.6, 3.25],
    )

    add_heading(doc, "Appendix A: Reproducible Command Set")
    add_code_block(
        doc,
        "Get-ChildItem -Force\n"
        "Get-Content composer.json\n"
        "Get-Content package.json\n"
        "Get-Content phpunit.xml\n"
        "Get-ChildItem -Path tests -Recurse -File | Select-Object FullName\n"
        "php -v\n"
        "php -m\n"
        "(Get-ChildItem -Path tests\\Unit -Recurse -Filter *.php | Measure-Object).Count\n"
        "(Get-ChildItem -Path tests\\Feature -Recurse -Filter *.php | Measure-Object).Count\n"
        "(Get-ChildItem -Path tests\\e2e -Recurse -Filter *.js | Measure-Object).Count\n"
        "$env:XDEBUG_MODE='coverage'; php artisan test --coverage\n"
        "php artisan test --compact\n"
        "php artisan test --testsuite=Unit --compact\n"
        "php artisan test --testsuite=Feature --compact\n"
        "Get-Date -Format \"yyyy-MM-dd HH:mm:ss zzz\"\n"
        "git rev-parse --short HEAD\n"
        "git status --short"
    )

    add_heading(doc, "Appendix B: Decision-Ready Statement")
    add_body(
        doc,
        "Based on the investigation, the defensible statement for a formal project report is: "
        '"The NUBL Laravel application has 93.3% PHPUnit test coverage over the app/ source tree. '
        'The PHP test suite includes both Unit and Feature tests: 243 Unit tests and 541 Feature tests, '
        'with all 784 tests passing across 2,658 assertions. Playwright E2E tests are also present as a separate '
        'browser testing layer, but they are not included in the PHPUnit coverage percentage."'
    )

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("NUBL Test Coverage Investigation Report")
    footer_run.font.name = "Aptos"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(COLORS["gray"])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
